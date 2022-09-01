from dataclasses import asdict, dataclass
from docx.shared import Cm


class BaseData:
    # @property
    @classmethod
    def data(cls):
        return {
            'Table1': {
                'class': Connection,
                'widths': [1.49, 4.75, 1.75, 8.49],
                'name': 'Тепловая нагрузка перспективного потребителя, '
                'источник тепловой энергии и ТСО, участвующие в подключении',
            },
            'table2': {
                'class': Event,
                'widths': [1.24, 5.50, 1.75, 2.50, 5.49],
                'name': 'Основные мероприятия и объемы капитальных затрат, '
                'необходиые для рассматриваемого подключения',
            },
            'Table3': {
                'class': TSO,
                'widths': [1.5, 8.0, 1.75, 1.75, 1.75, 1.75],
                'name': 'Расчет изменения НВВ после предлагаемого подключения',
            },
        }

    @classmethod
    def get_class(cls, table_name):
        return cls.data().get(table_name)['class']

    @classmethod
    def get_name(cls, table_name):
        return cls.data().get(table_name)['name']

    @classmethod
    def get_widths(cls, table_name):
        return cls.data().get(table_name)['widths']


class TableData:
    def __init__(self, wb, j, table_name):
        self.name = BaseData.get_name(table_name)
        self.widths = BaseData.get_widths(table_name)

        rng = wb.defined_names.get(table_name,
                                   scope=wb.sheetnames.index(str(j)))
        rng_dict = dict(rng.destinations)
        dest = rng_dict[str(j)]
        private_range = wb[str(j)][dest]

        data = []
        for row in private_range:
            cell_values = []
            for cell in row:
                cell_values.append(cell.value)
            data.append(BaseData.get_class(table_name)(*cell_values))

        self.data = data


@dataclass
class Connection:
    id: str = '№ п/п'
    title: str = 'Наименование мероприятия'
    units: str = 'Ед. изм.'
    value: str = 'Значения показателя'

    def __post_init__(self):
        self.id = str(self.id)
        self.title = str(self.title)
        self.units = str(self.units)

        if self.value is None:
            self.value = ''
        elif isinstance(self.value, float) or isinstance(self.value, int):
            self.value = format(
                self.value, "6,.2f"
            ).replace(",", " ").replace(".", ",")
        else:
            self.value = str(self.value)


@dataclass
class Event:
    id: str = '№ п/п'
    title: str = 'Наименование мероприятия'
    diameter: str = 'Диаметр, мм'
    length: str = 'Протяженность, м'
    capex: str = 'Капитальные затраты в ценах 2021 года, млн руб. без НДС'

    def __post_init__(self):
        self.id = str(self.id)
        self.title = str(self.title)

        if self.diameter is None:
            self.diameter = ''
        else:
            self.diameter = str(self.diameter)

        if isinstance(self.capex, float) or isinstance(self.capex, int):
            self.capex = format(
                self.capex, "6,.2f"
            ).replace(",", " ").replace(".", ",")
        else:
            self.capex = str(self.capex)

        if isinstance(self.length, float) or isinstance(self.length, int):
            self.length = format(
                self.length, "6,.1f"
            ).replace(",", " ").replace(".", ",")
        else:
            self.length = str(self.length)


@dataclass
class TSO:
    id: str = '№ п/п'
    title: str = 'Наименование показателя'
    units: str = 'Ед. изм.'
    old_nvv: str = 'НВВ'
    delta_nvv: str = 'Изменение НВВ'
    new_nvv: str = 'НВВ после мероприятий'

    def __post_init__(self):
        self.id = str(self.id)
        self.title = str(self.title)

        if self.units is None:
            self.units = ''
        else:
            self.units = str(self.units)

        if self.old_nvv is None:
            self.old_nvv = ''
        elif isinstance(self.old_nvv, float) or isinstance(self.old_nvv, int):
            self.old_nvv = format(
                self.old_nvv, "6,.2f"
            ).replace(",", " ").replace(".", ",")
        else:
            self.old_nvv = str(self.old_nvv)

        if self.delta_nvv is None:
            self.delta_nvv = ''
        elif isinstance(self.delta_nvv, float) or isinstance(self.delta_nvv,
                                                             int):
            self.delta_nvv = format(
                self.delta_nvv, "6,.2f"
            ).replace(",", " ").replace(".", ",")
        else:
            self.delta_nvv = str(self.delta_nvv)

        if self.new_nvv is None:
            self.new_nvv = ''
        elif isinstance(self.new_nvv, float) or isinstance(self.new_nvv, int):
            self.new_nvv = format(
                self.new_nvv, "6,.2f"
            ).replace(",", " ").replace(".", ",")
        else:
            self.new_nvv = str(self.new_nvv)


@dataclass
class Style:
    txt_style: str = '_Обычный'
    table_style: str = 'Table Grid'
    table_txt_style: str = '_Обычный_табл_10пт_по центру'
    table_name_style: str = '_Подпись таблицы'


class Table:
    def __init__(self, data, name, widths, number=1, appendix='',
                 style=Style()):
        self.data = data
        self.style = style
        self.cols_number = len(asdict(data[0]).values())
        self.rows_number = len(data)
        self.widths = widths

        if self.widths is not None:
            self.widths = [Cm(width) for width in self.widths]

        if name == '':
            self.table_name = f'Таблица {appendix}{number}'
        else:
            self.table_name = (
                f'Таблица {appendix}{number} - '
                f'{name}'
            )

    def __set_col_widths(self, table):
        for row in table.rows:
            for idx, width in enumerate(self.widths):
                row.cells[idx].width = width
                row.cells[idx].paragraphs[0].style = (
                    self.style.table_txt_style
                )

    def create_table(self, mydoc):
        mydoc.add_paragraph('', style=self.style.txt_style)
        mydoc.add_paragraph(self.table_name, style=self.style.table_name_style)

        table = mydoc.add_table(rows=self.rows_number, cols=self.cols_number)

        for row in range(self.rows_number):
            row_data = asdict(self.data[row]).values()
            for key, value in enumerate(row_data):
                table.cell(row, key).paragraphs[0].add_run(value)
        table.style = self.style.table_style

        if self.widths is not None:
            table.autofit = False
            self.__set_col_widths(table)

        mydoc.add_paragraph('', style=self.style.txt_style)

    def __str__(self):
        return self.table_name
