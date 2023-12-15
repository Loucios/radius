from classes import Style, Table, TableData
from docx import Document
from openpyxl import load_workbook
from progress.bar import Bar


def create_block(mydoc, wb, j, table_number=1, appendix='', style=Style()):
    '''Создает раздел книги

    Структура раздела следующая.
    1. Заголовок
    2. Абзац
    3. Таблица (с наименованием)
    4. Абзац
    5. Таблица (с наименованием)
    6. Абзац
    7. Таблица (с наименованием)
    '''
    ws = wb[str(j)]
    # Добавляем заголовок
    length = len(mydoc.paragraphs)
    mydoc.paragraphs[length - 1].style = '_1.'
    mydoc.paragraphs[length - 1].add_run(ws['D6'].value)

    #########################################################
    # Добываем данные из нужной таблицы файла Excel
    # (Книга, Лист, Наименование таблицы)
    connections = TableData(wb, j, 'Table1')
    # Добавляем абзац перед нашей таблицей
    mydoc.add_paragraph(
        f'В настоящем разделе рассматривается целесообразность '
        f'подключения к источнику тепловой энергии {connections.data[3].value}'
        f' следующей территории: '
        # f'{ws["D7"].value}: '
        f'{ws["D6"].value}. '
        f'В таблице {appendix}{table_number} приведены показатели '
        f'тепловой нагрузки рассматриваемого потребителя, а также '
        f'наименования ТСО, участвующих в подключении. '
        f'Приведен вывод о целесообразности рассматриваемоего подключения '
        f'на основе выполненных расчетов.',
        style=style.txt_style
    )
    # Добавляемы таблицу с заголовком и форматируем
    connections_table = Table(connections.data,
                              connections.name,
                              connections.widths,
                              table_number,
                              appendix)
    connections_table.create_table(mydoc)
    table_number += 1

    #########################################################
    # Добываем данные из нужной таблицы файла Excel
    # (Книга, Лист, Наименование таблицы)
    events = TableData(wb, j, 'table2')
    # Добавляем абзац перед нашей таблицей
    mydoc.add_paragraph(
        f'Произведена оценка необходимых капитальных затрат '
        f'для подключения рассматриваемоего потребителя к источнику '
        f'тепловой энергии {connections.data[3].value} '
        f'(таблица {appendix}{table_number}).',
        style=style.txt_style
    )
    # Добавляемы таблицу с заголовком и форматируем
    events_table = Table(events.data,
                         events.name,
                         events.widths,
                         table_number,
                         appendix)
    events_table.create_table(mydoc)
    table_number += 1

    #########################################################
    # Добываем данные из нужной таблицы файла Excel
    # (Книга, Лист, Наименование таблицы)
    tsos = TableData(wb, j, 'Table3')
    # Добавляем абзац перед нашей таблицей
    mydoc.add_paragraph(
        f'Произведен расчет изменения НВВ с целью определения '
        f'целесобразности подключения рассматриваемой территории '
        f'(таблица {appendix}{table_number}).',
        style=style.txt_style
    )
    # Добавляемы таблицу с заголовком и форматируем
    events_table = Table(tsos.data,
                         tsos.name,
                         tsos.widths,
                         table_number,
                         appendix)
    events_table.create_table(mydoc)
    table_number += 1

    # Запоминаем номер последней таблицы
    return table_number


def main():
    print('Загружаем Excel')
    wb = load_workbook(filename='RET5.xlsm', data_only=True)
    chapters_number = wb['Результат']['A1'].value

    books_number = 1
    appendix_number = 'Д'
    table_number = 1
    bar = Bar('Создаем Word', max=chapters_number)  # Индикатор выполнения
    for j in range(1, chapters_number + 1):
        # Разбиваем на книги
        if j % 125 == 0 or j == 1:
            mydoc = Document('my_doc.docx')
        # Создаем повторяющийся блок документа
        table_number = create_block(mydoc, wb, j, table_number,
                                    appendix_number)
        # Разбиваем на книги
        if j % 124 == 0 or j == chapters_number:
            mydoc.save(
                f'Книга 7 Глава 7 Приложение {appendix_number} том '
                f'{books_number}.docx'
            )
            books_number += 1
            table_number = 1
        bar.next()
    bar.finish()


if __name__ == '__main__':
    main()
